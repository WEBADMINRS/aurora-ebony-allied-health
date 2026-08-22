from docx import Document
from openpyxl import Workbook
from pathlib import Path

base = Path.home() / "Desktop" / "RS Assurance Group Resources"

def create_doc(path, title, sections):
    doc = Document()

    header = doc.sections[0].header
    header.paragraphs[0].text = "RS Assurance Group | Compliance Resources"

    doc.add_heading(title, level=1)

    doc.add_paragraph(
        "Document Owner: RS Assurance Group\n"
        "Version: 1.0\n"
        "Review Date: __________"
    )

    for heading, text in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(text)

    footer = doc.sections[0].footer
    footer.paragraphs[0].text = "RS Assurance Group | Professional Compliance Resources"

    doc.save(path)


incident_path = base / "Incident Management Pack"


create_doc(
    incident_path / "01 Templates" / "Incident Report Form.docx",
    "Incident Report Form",
    [
        ("Purpose",
        "This form supports consistent recording of incidents and relevant details."),
        ("Incident Details",
        "Incident date:\nLocation:\nPeople involved:\nDescription of incident:"),
        ("Immediate Actions",
        "Actions taken immediately following the incident:"),
        ("Review and Follow Up",
        "Investigation outcome, corrective actions and follow up requirements:")
    ]
)


create_doc(
    incident_path / "01 Templates" / "Incident Investigation Template.docx",
    "Incident Investigation Template",
    [
        ("Purpose",
        "A structured template to review incidents and identify contributing factors."),
        ("Investigation Details",
        "Investigator:\nDate reviewed:\nIncident reference:"),
        ("Findings",
        "Summary of findings and contributing factors:"),
        ("Recommendations",
        "Actions recommended to prevent recurrence:")
    ]
)


create_doc(
    incident_path / "01 Templates" / "Corrective Action Record.docx",
    "Corrective Action Record",
    [
        ("Issue Identified",
        "Describe the issue requiring corrective action."),
        ("Action Required",
        "Outline actions required."),
        ("Completion Tracking",
        "Responsible person:\nDue date:\nCompletion date:")
    ]
)


wb = Workbook()

ws = wb.active
ws.title = "Incident Register"

ws.append([
    "Incident ID",
    "Date",
    "Description",
    "Severity",
    "Actions Taken",
    "Status",
    "Review Date"
])

wb.save(
    incident_path / "02 Registers" / "Incident Register.xlsx"
)


create_doc(
    incident_path / "03 Guidance" / "Incident Management Procedure.docx",
    "Incident Management Procedure",
    [
        ("Purpose",
        "This procedure outlines a structured approach for recording, reviewing and managing incidents."),
        ("Recording Requirements",
        "Incidents should be documented consistently with sufficient detail."),
        ("Review Process",
        "Incidents should be reviewed to identify improvements and preventative actions.")
    ]
)

print("Incident Management Pack created successfully.")

